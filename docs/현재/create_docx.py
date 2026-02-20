"""
관광벤처 예비부문 사업계획서 DOCX 생성 스크립트 v4
- BM 개편: ₩3,900/월 단일 플랜, 일상 콘텐츠 앱화
- 출력: 관광벤처_사업계획서_Fryndo_v4.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 기본 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_font(cell, text, bold=False, size=10, alignment=None):
    """셀에 텍스트와 서식 설정"""
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    # 셀 내부 여백
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, color="D9E2F3"):
    """셀 배경색"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(text, level=1):
    """제목 추가"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Malgun Gothic'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return h


def add_para(text, bold=False, size=10, space_after=6):
    """문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p.paragraph_format.space_after = Pt(space_after)
    return p


def create_table(headers, rows, header_color="4472C4", col_widths=None):
    """표 생성 헬퍼 (진한 헤더, 흰색 글자)"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, header_color)
        # 흰색 글자
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 데이터
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_font(cell, str(val), size=9)

    # 열 너비
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 표 아래 간격
    return table


def create_table_light(headers, rows, first_col_bold=True):
    """밝은 스타일 표 (헤더: 밝은 파란 배경, 검은 글자)"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "D9E2F3")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            is_bold = first_col_bold and c_idx == 0
            set_cell_font(cell, str(val), bold=is_bold, size=9)

    doc.add_paragraph()
    return table


# ===================================================
# 문서 시작
# ===================================================

# ── 표지 ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('예비관광벤처부문 사업계획서')
run.font.size = Pt(26)
run.font.bold = True
run.font.name = 'Malgun Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Fryndo (프린도)')
run.font.size = Pt(18)
run.font.name = 'Malgun Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('혼자 여행하는 관광객을 위한 AI 기반 동반자 매칭 및\n지역 관광지 디지털 수집(NFT) 플랫폼')
run.font.size = Pt(12)
run.font.name = 'Malgun Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026년 2월')
run.font.size = Pt(14)
run.font.name = 'Malgun Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name_p.add_run('김다운')
run.font.size = Pt(14)
run.font.name = 'Malgun Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_page_break()


# ===================================================
# 일반현황
# ===================================================

add_heading_styled('일반현황', level=1)

create_table_light(
    ['항목', '기입'],
    [
        ['대표자 성명', '김다운'],
        ['생년월일', '1988.11.22'],
        ['성별', '남'],
        ['기업(예정)명', '주식회사 프린도'],
        ['본사 소재지', '경기도 광주시'],
        ['사업자 번호', '(미등록 - 예비창업)'],
        ['개업연월일', '(예정)'],
        ['신청유형', '관광체험서비스'],
        ['서비스(상품)명', 'Fryndo (프린도)'],
        ['한줄설명', '혼자 여행하는 관광객을 위한 AI 기반 동반자 매칭 및 지역 관광지 디지털 수집(NFT) 플랫폼'],
    ]
)


# ── 협약목표 ──
add_heading_styled('협약목표 (협약기간내)', level=2)

add_para('성과지표(KPI): 상품서비스 개시(2건), 아이템 검증(20회), 고객확보(5,000명), 홍보마케팅(25건), 모바일 어플리케이션 개발(1개)', bold=True)

create_table_light(
    ['항목', '목표'],
    [
        ['매출 (백만원)', '7'],
        ['관광매출 (%)', '100'],
        ['투자유치 (백만원)', '0'],
        ['신규고용 (명)', '4'],
    ]
)


# ── 3개년 주요성과 ──
add_heading_styled('3개년 주요성과', level=2)
add_para('※ 예비창업자로 해당 없음')

create_table_light(
    ['구분', '2024년', '2023년', '2022년'],
    [
        ['매출 (백만원)', '-', '-', '-'],
        ['관광매출 (%)', '-', '-', '-'],
        ['투자유치 (백만원)', '-', '-', '-'],
        ['신규고용 (명)', '-', '-', '-'],
    ]
)

doc.add_page_break()


# ===================================================
# 사업요약
# ===================================================

add_heading_styled('예비관광벤처 사업요약 (1매 이내)', level=1)

summary_data = [
    ['사업명', 'Fryndo (프린도) - AI 기반 관광 동반자 매칭 및 지역 관광지 NFT 수집 플랫폼'],
    ['사업개요', '1인 여행자 비율 25% 이상(한국관광공사)이나 67%가 불편함, 여성 78%가 안전 최대 우려(Booking.com). Fryndo는 ①AI 동반자 매칭(여행스타일·일정·관심사 분석, 실명인증, 실시간 위치공유) ②GPS 관광지 NFT 수집(지역 한정 디지털 수집품으로 소도시 방문 유도) ③관광 안전 케어(긴급SOS, 신뢰점수)로 1인 관광객의 동행·안전·기록 문제를 해결합니다.'],
    ['시장성', 'TAM 7조원(국내 관광 28조×1인 25%) → SAM 2.8조원(디지털 40%) → SOM 28억원(점유 0.1%). 글로벌 CAGR 14.3%. 수익: B2C 구독 ₩3,900/월(50%) + B2B 지자체 제휴(25%) + 광고(20%) + NFT(5%). 평소 콘텐츠·커뮤니티 + 여행 시 매칭으로 일상 앱화. 1차 타겟: 제주 25~34세 여성.'],
    ['사업화역량 및\n실현가능성', 'MVP 핵심 기능 프로토타입 구현(AI매칭, 채팅, GPS인증, NFT민팅). URL: https://various-belva-untab-1a59bee2.koyeb.app. 자금 7,000만원: 인건비 6,250만(기획PM+개발자2+마케터1) + 외주·광고·인프라 등 750만. M3 앱 출시 → 제주 시범 → 확대.'],
    ['관광산업 트렌드\n부합성 및 연관성', '①1인 관광 확대→동행 매칭 ②K-관광 디지털 전환→AI+블록체인 ③지역 분산 정책→지역 NFT ④MZ 경험 소비→게이미피케이션 ⑤관광 안전→위치공유·SOS ⑥데이터 활용→여행패턴 수집'],
    ['팀빌딩 및 역량', '대표 김다운: 9년차 풀스택(Spring Boot, React, Web3). 크리에이티브힐 NFT, 에너지마켓플레이스 매칭 알고리즘, IBK 다중인증 경험. MVP 단독 개발로 실행력 입증. 선정 후 기획PM(전담)+앱 개발자+블록체인 개발자+마케터 4명 채용, 디자인 외주. 대표는 기술 총괄(백엔드/AI).'],
    ['기업가 정신\n및 태도', '제주도 1년 거주(개발자 근무) 경험에서 1인 관광객의 동행 니즈를 직접 목격. 동남아 배낭여행에서도 혼자 여행하는 한국인들의 불편함을 체감. 투자·팀원 없이 MVP 프로토타입을 단독 구현한 실행력으로 관광 동행 인프라를 만들겠다는 확신.'],
    ['이미지 (해당시)', 'MVP 스크린샷: ①랜딩페이지 ②로그인(카카오/구글) ③여행그룹 ④실시간채팅 ⑤NFT수집 ⑥AI매칭'],
]

table = doc.add_table(rows=len(summary_data), cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, content) in enumerate(summary_data):
    cell0 = table.rows[i].cells[0]
    cell1 = table.rows[i].cells[1]
    set_cell_font(cell0, label, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cell0, "D9E2F3")
    set_cell_font(cell1, content, size=9)
    cell0.width = Cm(3.5)
    cell1.width = Cm(13)

doc.add_paragraph()
doc.add_page_break()


# ===================================================
# 1. 사업 개요
# ===================================================

add_heading_styled('1. 사업 개요 (2장 이내)', level=1)
add_heading_styled('1-1. 문제 인식의 구체성', level=2)

add_para('(1) 사업 아이템이 해결하고자 하는 구체적인 문제', bold=True)
add_para('Fryndo는 1인 관광객이 겪는 세 가지 핵심 문제를 해결합니다.')

# 문제 3가지
def add_bullet_item(title_text, body_text):
    """굵은 제목 + 일반 본문의 단락"""
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.font.bold = True
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(10)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    run2 = p.add_run(body_text)
    run2.font.name = 'Malgun Gothic'
    run2.font.size = Pt(10)
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_bullet_item(
    '■ 동행 부재: ',
    '1인 여행 시 불편함 경험 67%, "동행 있으면 좋겠다" 응답 72%. 혼밥 어색함, 사진 부탁할 사람 없음, 현지 교류 한계.'
)
add_bullet_item(
    '■ 안전 불안: ',
    '여성 1인 여행자 78%가 "안전이 최대 걱정"(Booking.com, 2025). 야간 이동 불안, 응급상황 대처 곤란, 안전 우려로 여행 취소 34%.'
)
add_bullet_item(
    '■ 여행 기록 휘발: ',
    '여행 후 기록 남기는 비율 23%. 방문 인증 체계 부재, "여행을 수집하고 싶다" 응답 71%.'
)

doc.add_paragraph()

add_para('(2) 문제 발생의 주요 원인 및 배경', bold=True)
add_para('• 1인 여행 비율 25% 이상(한국관광공사, 2024), 아시아 1인 여행 검색 +16% YoY(Agoda, 2025), 한국인 80%가 2026년 더 많이 여행 계획(Skyscanner, 2026)')
add_para('• 1인 여행 수요는 급증하지만, 동행을 찾는 인프라는 전무. 기존 플랫폼(트립어드바이저, 에어비앤비)은 숙소·관광지 정보에 집중, 동반자 매칭 기능 없음.')
add_para('• 서울/경기 관광객 65% 집중, 비수도권 관광 수입 감소 → 지역 관광 활성화 수단 부족')

add_para('(3) 문제의 영향을 받는 구체적인 대상', bold=True)

create_table_light(
    ['대상', '규모', '핵심 니즈'],
    [
        ['25~34세 여성 1인 여행자', '아태 1인 여행자의 30%', '안전한 동행, 사진 촬영, 맛집 공유'],
        ['MZ세대 전체 (25~44세)', '53%', '경험 소비, 디지털 수집, 커뮤니티'],
        ['비수도권 지자체·관광공사', '전국 17개 시도', '관광객 유치, 지역 콘텐츠 부족'],
        ['외국인 1인 관광객', 'K-관광 수요 증가', '현지인/한국인 교류, 언어 장벽'],
    ]
)

add_para('(4) 문제 해결 방법의 현실성 및 차별성', bold=True)

create_table_light(
    ['기존 서비스', '한계', 'Fryndo 차별점'],
    [
        ['트립어드바이저', '동반자 매칭 없음', '여행 스타일 기반 AI 동행 매칭'],
        ['에어비앤비 익스피리언스', '일회성 투어, 지속적 관계 X', '그룹 단위 장기 동행'],
        ['Couchsurfing', '신뢰도·안전 이슈', '실명인증, 신뢰점수, 긴급SOS'],
        ['틴더 여행모드', '로맨스 중심, 여행 특화 X', '순수 여행 동행 목적'],
    ]
)

add_para('• 현실성: MVP 프로토타입 구현(실제 서비스 접속 가능), 9년 개발 경력으로 기술 실행력 입증')
add_para('• 차별성: 동행 매칭 + GPS 관광지 NFT 수집 + 안전 시스템을 하나의 플랫폼에서 제공')

doc.add_page_break()


# ===================================================
# 2. 시장성
# ===================================================

add_heading_styled('2. 시장성 (3장 이내)', level=1)
add_heading_styled('2-1. 시장현황 및 목표고객', level=2)

add_para('(1) 시장 현황 및 향후 전망', bold=True)

create_table_light(
    ['구분', '규모', '산출 근거'],
    [
        ['TAM', '7조원', '국내 관광시장 28조 × 1인 여행자 25%'],
        ['SAM', '2.8조원', 'TAM × 디지털 플랫폼 이용률 40%'],
        ['SOM', '28억원', 'SAM × 초기 점유율 0.1%'],
    ]
)

add_para('• 글로벌 1인 여행 시장: 2024년 $4,823억 → 2030년 $1.07조 (CAGR 14.3%, Grand View Research)')
add_para('• 아시아태평양 CAGR 16.1%로 가장 빠른 성장 지역')
add_para('• 한국인 41%가 여행에서 친구/로맨스 추구(Skyscanner, 2026)')

add_para('(2) 목표고객 및 제공 가치', bold=True)

create_table_light(
    ['구분', '타겟', '제공 가치'],
    [
        ['1차 (초기 집중)', '제주도 25~34세 여성 1인 여행자', '안전한 동행 매칭, 실시간 위치 공유, 여행 사진 공유'],
        ['2차 (확대)', 'MZ세대 전체, 부산/강릉/경주', '여행 게이미피케이션(NFT 수집), 커뮤니티'],
        ['3차 (글로벌)', '외국인 관광객', 'K-관광 현지 동행, 다국어 지원'],
    ]
)


add_heading_styled('2-2. 사업모델(Business Model)', level=2)

# BM 도식화를 표로
add_para('■ 핵심 전략: 평소에는 콘텐츠·커뮤니티, 여행 시에는 매칭·인증 → 일상적 사용 유도', bold=True)

create_table_light(
    ['구분', '무료', 'Plus ₩3,900/월'],
    [
        ['평소 (매일)', '여행 콘텐츠 피드, NFT 갤러리, 커뮤니티', '+ AI 맞춤 추천, 프리미엄 콘텐츠'],
        ['여행 시', '기본 매칭 월 3회, Common NFT', '+ 무제한 매칭, Rare NFT 부스트'],
    ]
)

add_para('■ 수익 구조', bold=True)

create_table_light(
    ['수익원', '비중', '내용'],
    [
        ['B2C 구독', '50%', 'Plus ₩3,900/월 (커피 한 잔 가격 → 높은 전환율 8~10%)'],
        ['B2B 관광기관 제휴', '25%', '지자체 NFT 발행 대행, 관광 데이터 제공'],
        ['광고', '20%', '관광지·숙박·맛집 타겟 광고 (DAU 기반)'],
        ['NFT/기타', '5%', 'NFT 거래 수수료'],
    ]
)

add_para('• ₩3,900/월 = 부담 없는 구독 → 높은 전환율(8~10%) 기대')
add_para('• LTV:CAC = 6:1 (CAC ₩8,000 / LTV ₩46,800, 평균 12개월 유지)')
add_para('• NFT 의존도 5%로 최소화, 구독+B2B+광고로 매출 95% 유지')


add_heading_styled('2-3. 잠재리스크 및 대응방안', level=2)

create_table_light(
    ['잠재 리스크', '대응방안'],
    [
        ['동행 간 안전사고·분쟁', '실명인증(신분증/SNS 연동), 후기·평점 시스템, 긴급 SOS, 실시간 위치 공유, 신뢰 점수제'],
        ['NFT 시장 침체', '투기 아닌 "수집/인증" 목적 포지셔닝, 매출 의존도 5%로 최소화, 구독+B2B 중심 수익'],
        ['초기 사용자 확보 어려움', '제주도 1개 지역 집중 → PMF 검증 후 확대, 인플루언서 10명 시딩, 커뮤니티 바이럴'],
        ['유료 전환율 저조', '₩3,900 저가 전략 + 평소 콘텐츠 락인으로 전환 유도, 무료체험 기간 제공'],
        ['경쟁 서비스 출현', 'MVP 선점 우위, 커뮤니티 락인, 빠른 개발 사이클'],
    ]
)

doc.add_page_break()


# ===================================================
# 3. 사업화 역량 및 실현 가능성
# ===================================================

add_heading_styled('3. 사업화 역량 및 실현 가능성 (4장 이내)', level=1)
add_heading_styled('3-1. 상품·서비스·인프라 개발 계획', level=2)

add_para('MVP(최소기능제품) 핵심 기능 프로토타입 구현', bold=True)

# 기능 상태 표 - 9행 (모바일앱 포함)
func_table = doc.add_table(rows=10, cols=3)
func_table.style = 'Table Grid'
func_table.alignment = WD_TABLE_ALIGNMENT.CENTER

func_headers = ['기능', '상태', '적용 기술']
for i, h in enumerate(func_headers):
    cell = func_table.rows[0].cells[i]
    set_cell_font(cell, h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cell, "D9E2F3")

func_data = [
    ['회원가입/로그인', '구현', 'Spring Security, JWT, OAuth 2.0 (카카오/구글)'],
    ['프로필 관리', '구현', 'PostgreSQL'],
    ['여행 그룹 생성/참여', '구현', 'Spring Boot REST API'],
    ['실시간 채팅', '구현', 'WebSocket, STOMP'],
    ['AI 매칭 추천', '개선중', 'OpenAI API (정확도 고도화 필요)'],
    ['GPS 위치 인증', '구현', 'Geolocation API'],
    ['NFT 민팅', '개선중', 'Polygon, Web3.js, IPFS (테스트넷 단계)'],
    ['반응형 웹', '구현', 'React, Tailwind CSS'],
    ['모바일 앱', '미착수', 'React Native (선정 후 개발)'],
]

for r_idx, row_data in enumerate(func_data):
    for c_idx, val in enumerate(row_data):
        cell = func_table.rows[r_idx + 1].cells[c_idx]
        align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 1 else None
        set_cell_font(cell, val, size=9, alignment=align)

doc.add_paragraph()

add_para('• 서비스 URL: https://various-belva-untab-1a59bee2.koyeb.app')
add_para('• GitHub: https://github.com/araeLaver/down')
add_para('• 향후 개발 계획: 모바일 앱 개발(React Native, M1~M3) → 정식 출시(M3) → AI 매칭 고도화 → 관광지 NFT 확대')


add_heading_styled('3-2. 홍보 및 판로개척', level=2)

create_table_light(
    ['채널', '비중', '실행 내용', '시기'],
    [
        ['인스타그램', '45%', '제주 여행 릴스, 여성 마이크로 인플루언서 10명 시딩', 'M3~'],
        ['여행 커뮤니티', '25%', '네이버 카페(혼자여행클럽), 에브리타임 바이럴', 'M3~'],
        ['콘텐츠 마케팅', '20%', '유튜브 숏폼(동행 후기), 블로그 SEO', 'M4~'],
        ['PR/제휴', '10%', '제주관광공사 채널 연계, 지역 매체 보도', 'M5~'],
    ]
)

add_para('• 판로: B2C 앱스토어 출시 + B2B 제주관광공사 MOU(NFT 발행 대행) + B2G 관광 데이터 제공')
add_para('• 바이럴 루프: 여행 콘텐츠 소비 → 동행 매칭 → 여행 후기·NFT 공유 → 친구 초대(Rare NFT 보상) → 신규 가입')
add_para('• 리텐션 전략: 평소에도 여행 콘텐츠·커뮤니티로 DAU 유지 → 여행 시즌 자연스러운 유료 전환')


add_heading_styled('3-3. 자금운용계획', level=2)

add_para('(1) 소요자금 및 조달 계획', bold=True)
add_para('• 법인 설립 후 사업 개시. 초기 자본금은 대표 자기자금으로 조달.')
add_para('• 정부지원금(관광벤처 사업화지원금) 7,000만원 + 자부담(현물 10%) 700만원 = 총 7,700만원')

add_para('(2) 정부지원금 집행계획', bold=True)
add_para('[표1] 자금집행계획 (정부지원금 7,000만원 기준)', bold=True, size=11)

# [표1] - 예산표 (10행 데이터 + 1행 헤더 = 11행)
table1 = doc.add_table(rows=11, cols=4)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

t1_headers = ['조달방법', '용도', '산출근거', '금액(천원)']
t1_data = [
    ['정부지원금', '기획/PM 인건비', '월 500만원 × 6개월 (전담, 사업운영·일정관리·제휴)', '30,000'],
    ['정부지원금', '앱 개발자 인건비', '월 450만원 × 3개월 (계약직, React Native 경력 2년+)', '13,500'],
    ['정부지원금', '블록체인 개발자 인건비', '월 500만원 × 2개월 (프리랜서, Solidity/Web3)', '10,000'],
    ['정부지원금', '마케터 인건비', '월 300만원 × 3개월 (관광/여행 경력)', '9,000'],
    ['정부지원금', 'UI/UX 디자인 외주', '앱 디자인 일괄 (프리랜서)', '2,000'],
    ['정부지원금', '광고선전비', '인스타그램 광고 + 인플루언서 시딩', '2,000'],
    ['정부지원금', '서버/API 인프라', '클라우드, AI API, 블록체인 가스비', '1,500'],
    ['정부지원금', '사무실·기타', '공유오피스, 법인설립, 세무기장, 상표 출원', '2,000'],
    ['', '합계', '', '70,000'],
    ['자기자금', '자부담 (현물 10%)', '대표 인건비 + 자택 사무실', '7,000'],
]

for i, h in enumerate(t1_headers):
    cell = table1.rows[0].cells[i]
    set_cell_font(cell, h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cell, "4472C4")
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for r_idx, row_data in enumerate(t1_data):
    for c_idx, val in enumerate(row_data):
        cell = table1.rows[r_idx + 1].cells[c_idx]
        align = WD_ALIGN_PARAGRAPH.RIGHT if c_idx == 3 else None
        is_bold = (r_idx == 8)  # '합계' row
        set_cell_font(cell, val, bold=is_bold, size=9, alignment=align)

doc.add_paragraph()
add_para('※ 대표 인건비는 정부지원금 집행 불가 → 자부담(현물)으로 처리')

doc.add_page_break()


add_heading_styled('3-4. 일정계획', level=2)

add_para('[표2] 사업추진 일정 (전체 사업단계)', bold=True, size=11)

create_table_light(
    ['추진기간', '추진내용', '세부내용'],
    [
        ['2026.05~07', '서비스 개발', '모바일 앱 개발(React Native), UI/UX 디자인, NFT 시스템 고도화'],
        ['2026.06~07', '시범운영', '제주도 베타 테스트, 테스터 200명 모집, 사용자 인터뷰'],
        ['2026.07', '정식 출시', '앱스토어/구글플레이 출시, 여름 성수기 마케팅'],
        ['2026.08~09', '성장', '부산/강릉 확대, B2B 지자체 제휴, NFT 컬렉션 확대'],
        ['2026.10~11', '성과 창출', '가을 관광 시즌 공략, 성과 정리, IR 준비'],
        ['2027.01~', '투자유치', 'Seed 라운드 추진, 전국 확대, 외국인 관광객 타겟'],
        ['2028~', '글로벌', '일본/태국 진출, 다국어 지원'],
    ]
)


add_para('[표3] 사업추진 일정 (협약기간 내, 2026.5~11월)', bold=True, size=11)

# 간트 차트 표 (5행 데이터 + 1행 헤더 + 1행 범례 = 7행)
gantt = doc.add_table(rows=7, cols=8)
gantt.style = 'Table Grid'
gantt.alignment = WD_TABLE_ALIGNMENT.CENTER

gantt_headers = ['추진내용', '5월', '6월', '7월', '8월', '9월', '10월', '11월']
for i, h in enumerate(gantt_headers):
    cell = gantt.rows[0].cells[i]
    set_cell_font(cell, h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cell, "D9E2F3")

gantt_data = [
    ['상품·서비스 개발', '■', '■', '■', '□', '□', '', ''],
    ['홍보·판로개척', '', '', '□', '■', '■', '■', '□'],
    ['창업·업종 추가 계획', '■', '', '', '', '', '', ''],
    ['인력·조직 수립, 개발', '■', '■', '■', '■', '□', '□', ''],
]

for r_idx, row_data in enumerate(gantt_data):
    for c_idx, val in enumerate(row_data):
        cell = gantt.rows[r_idx + 1].cells[c_idx]
        align = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else None
        set_cell_font(cell, val, size=9, alignment=align)
        # 색칠
        if val == '■':
            shade_cell(cell, "4472C4")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif val == '□':
            shade_cell(cell, "B4C6E7")

# 범례 행 (마지막 행, row index 5)
# 먼저 빈 데이터 행 index 5 처리
for c_idx in range(8):
    cell = gantt.rows[5].cells[c_idx]
    set_cell_font(cell, '', size=8)

# KPI 행 (row index 6)
kpi_row = gantt.rows[6]
kpi_row.cells[0].merge(kpi_row.cells[7])
set_cell_font(kpi_row.cells[0], '■ = 집중 추진, □ = 유지/운영', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# 협약 기간 내 목표를 간트 표 아래 5행(범례 바로 위)에 merge로 표시하는 대신
# 별도 텍스트로 표 아래에 표시
add_para('협약기간 내 목표: 상품서비스 개시(2건), 아이템 검증(20회), 고객확보(5,000명), 매출 7백만원, 투자유치 0백만원, 신규고용 4명', bold=True)


add_heading_styled('3-5. 성장전략 (Scale-up 계획)', level=2)

add_para('사업화 이후 3년간 성장 계획. 초기 제주 집중 → 전국 확대 → 글로벌 진출.')
add_para('• 손익분기점: 유료 구독자 5,000명(₩3,900×5,000=월 1,950만), MAU 60,000명 도달 시 (2028년 Q2 예상)')

add_para('[표4] Scale-up 계획', bold=True, size=11)

create_table_light(
    ['구분', '2026년', '2027년', '2028년', '근거'],
    [
        ['1. 매출액(백만원)', '7', '120', '600', '구독(₩3,900)+B2B+광고. 저가 구독으로 빠른 확산'],
        ['2. 투자유치(백만원)', '0', '300', '1,000', '2027 Seed, 2028 Series A'],
        ['3. 신규고용(명)', '4', '7', '12', '기획PM+개발+마케팅+CS 순차 확대'],
        ['4. 제휴기관(개)', '2', '10', '30', '지자체/관광공사/여행사'],
        ['5. 서비스이용자(명)', '5,000', '30,000', '100,000', 'MAU 기준, 제주→전국→해외'],
    ]
)

doc.add_page_break()


# ===================================================
# 4. 관광산업 트렌드
# ===================================================

add_heading_styled('4. 관광산업 트렌드 부합성 및 연관성 (1장 이내)', level=1)
add_heading_styled('4-1. 관광시장 이해도 및 트렌드 부합성', level=2)

create_table_light(
    ['관광 트렌드', '현황', 'Fryndo의 대응'],
    [
        ['1인 관광 시장 확대', '1인 여행 25%+, 연 16% 검색 증가', '신뢰 기반 동행 매칭으로 여행 장벽 해소'],
        ['K-관광 디지털 전환', '정부 스마트관광 정책 추진', 'AI 매칭 + 블록체인 인증 = 관광 테크 선도'],
        ['지역 관광 분산 정책', '서울/경기 65% 집중', '지역 한정 NFT로 소도시 방문 유도'],
        ['MZ세대 경험 소비', 'MZ세대 관광 시장 주도', '디지털 수집품 + 커뮤니티 = 게이미피케이션'],
        ['관광 안전 강화', '여성 안전 이슈 사회적 관심', '실시간 위치 공유, 실명인증, 긴급 SOS'],
        ['관광 데이터 활용', '공공 관광 데이터 구축 정책', '여행 패턴, 인기 관광지 데이터 수집·분석·제공'],
    ]
)


add_heading_styled('4-2. 사업 아이템과 관광산업과의 연관성', level=2)

add_para('Fryndo는 관광산업 자체가 사업 영역입니다. 숙소·티켓이 아닌 "관광 동행"이라는 새로운 관광 인프라를 제공합니다.')

add_para('1. 지역 관광 균형 발전', bold=True)
add_para('지역 한정 NFT 수집 시스템으로 관광객이 소도시를 방문해야만 획득할 수 있는 디지털 수집품 제공 → 자연스러운 지역 분산')

add_para('2. 관광 안전망 구축', bold=True)
add_para('1인 관광객(특히 여성·고령자)을 위한 동행 매칭과 안전 시스템으로 관광 안전 인프라 보완')

add_para('3. 스마트 관광 생태계', bold=True)
add_para('AI 매칭, 블록체인 인증, GPS 위치 서비스 등 첨단 기술을 관광에 접목하여 K-관광 디지털 경쟁력 강화')

add_para('4. 관광 빅데이터 구축', bold=True)
add_para('여행 패턴, 인기 관광지, 체류시간, 동행 선호도 데이터를 수집하여 관광공사·지자체에 제공 → 데이터 기반 관광 정책 수립 기여')

doc.add_page_break()


# ===================================================
# 5. 팀 빌딩 및 역량
# ===================================================

add_heading_styled('5. 팀 빌딩 및 역량 (2장 이내)', level=1)
add_heading_styled('5-1. 사업자의 전문성 및 운영 능력', level=2)

add_para('대표자: 김다운 | 9년차 풀스택 개발자 (2016~현재)', bold=True, size=11)

create_table_light(
    ['경험', '내용', 'Fryndo 직접 적용'],
    [
        ['블록체인/NFT', '크리에이티브힐 Polygon ERC-3525 NFT 개발', 'GPS 인증 NFT 민팅 시스템'],
        ['매칭 알고리즘', '에너지마켓플레이스 매칭 시스템 설계/개발', 'AI 동반자 매칭 시스템'],
        ['인증 시스템', 'IBK 기업은행 전자약정 다중인증 구현', '실명인증, 소셜 로그인'],
        ['위치 서비스', '서귀포시/고양시 지자체 위치기반 서비스', 'GPS 관광지 방문 인증'],
        ['팀 리딩', '헥토데이터 파트장 3년', '개발팀 관리, 프로젝트 수행'],
    ]
)

add_para('기술 스택: Java, Spring Boot, React, Python, NestJS, Polygon, Web3.js, PostgreSQL, AWS, Docker')
add_para('사업화 의지: 투자도 팀원도 없이 MVP 프로토타입을 혼자 기획·설계·개발. 핵심 기능(매칭·채팅·NFT)을 실제 동작하는 수준까지 구현하여 기술 실행력 입증. 사업 운영은 전담 기획/PM을 채용하여 위임하고, 대표는 기술 총괄(백엔드·AI·블록체인)에 집중합니다.', bold=True)


add_heading_styled('5-2. 팀원과의 신뢰 구축 능력', level=2)

add_para('• 팀 리딩 경험: 헥토데이터에서 3년간 파트장으로 개발팀을 이끈 경험. 업무 배분, 코드 리뷰, 일정 관리 등 팀 운영 역량 보유.')
add_para('• 다양한 협업 경험: 9년간 7개 회사에서 SI/서비스 프로젝트 수행. 기획자, 디자이너, 타 개발자와의 협업에 익숙.')
add_para('• 조직 문화 방향: 주간 스프린트 기반 애자일 운영, 매일 15분 스탠드업 미팅, 코드 리뷰 필수화. 소규모 팀에서 빠르게 실행하고 빠르게 피드백 반영.')
add_para('• 외부 네트워크: 법률(이용약관 검토 완료), 세무(기장 위탁 확보), 디자인(프리랜서 섭외 완료) 등 외부 전문가 협력 체계 구축.')


add_heading_styled('5-3. 인력 및 조직 운영계획', level=2)

add_para('경영원칙: "린 스타트업 + 빠른 실행". 핵심 기능에 집중하고 사용자 피드백 기반으로 빠르게 개선. PMF 검증 후 팀 확대.', bold=True)

add_para('[표5] 팀 구성 현황', bold=True, size=11)

# [표5] - 6행 데이터 (대표 + 기획PM + 앱개발 + 블록체인 + 마케터 + 디자이너)
team_table = doc.add_table(rows=7, cols=5)
team_table.style = 'Table Grid'
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER

team_headers = ['성명', '직위', '담당업무', '보유역량(경력/학력 등)', '구성상태']
for i, h in enumerate(team_headers):
    cell = team_table.rows[0].cells[i]
    set_cell_font(cell, h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cell, "D9E2F3")

team_data = [
    ['김다운', '대표', '기술 총괄\n(백엔드·AI·블록체인)', '9년차 풀스택, Web3/NFT, 매칭 알고리즘\n방송통신대 컴공 재학', '완료'],
    ['(채용예정)', '기획/PM', '사업운영 전담, 일정관리,\n외부 제휴, 서비스 기획', '스타트업/IT PM 경력 2년+\n월 500만', '예정\n(26.05)'],
    ['(채용예정)', '개발자', 'React Native\n모바일 앱 개발', 'React Native 경력 2년+\n월 450만', '예정\n(26.05)'],
    ['(채용예정)', '개발자', 'NFT 스마트컨트랙트\nWeb3 연동', 'Solidity/Polygon 경험\n월 500만', '예정\n(26.06)'],
    ['(채용예정)', '마케터', 'SNS, 인플루언서\n커뮤니티 운영', '여행/관광 마케팅 경험\n월 300만', '예정\n(26.09)'],
    ['(외주)', '디자이너', '앱 UI/UX 디자인\n브랜딩', '모바일 앱 전문 프리랜서', '예정\n(26.05)'],
]

for r_idx, row_data in enumerate(team_data):
    for c_idx, val in enumerate(row_data):
        cell = team_table.rows[r_idx + 1].cells[c_idx]
        is_bold = (c_idx == 0)
        set_cell_font(cell, val, bold=is_bold, size=9)

doc.add_paragraph()
doc.add_page_break()


# ===================================================
# 6. 기업가 정신 및 태도
# ===================================================

add_heading_styled('6. 기업가 정신 및 태도 (1장 이내)', level=1)
add_heading_styled('6-1. KPI 달성의 현실성 및 사업 참여 적극성', level=2)

add_para('■ 왜 이 문제를 풀어야 하는가', bold=True, size=11)

add_para('저는 제주도에서 1년간 거주하며 개발자로 일한 경험이 있습니다. 오션스솔루션에서 JDC(제주국제자유도시개발센터) 그룹웨어를 개발하고, 서귀포시 웹사이트를 유지보수하면서 제주 지역 관광 생태계를 가까이에서 경험했습니다.')

add_para('제주에서 살면서 매일 마주친 것은 혼자 여행 온 관광객들이었습니다. 카페에서, 올레길에서, 게스트하우스에서 혼자 온 여행자들이 "같이 갈 사람 없나요?"라고 묻는 장면을 수없이 봤습니다. 1인 여행자가 가장 많이 찾는 관광지에서 1년을 살았기에, 이들이 겪는 불편함\u2015혼밥의 어색함, 사진 부탁할 사람 없음, 야간 이동 시 불안\u2015을 생활 속에서 체감할 수 있었습니다.')

add_para('이후 동남아 배낭여행을 다녀오면서도 같은 문제를 느꼈습니다. 혼자 여행하는 한국인들이 많았지만, 서로 만날 수 있는 체계적인 방법이 없었습니다. 제주에서의 1년, 그리고 해외 여행 경험이 합쳐져 "안전하고 신뢰할 수 있는 동행 매칭 플랫폼"이라는 아이디어가 구체화되었습니다.')

add_para('제주 지역 IT 기업에서 일하며 쌓은 지자체 시스템 개발 경험(서귀포시, JDC)은 향후 지자체·관광공사 B2B 제휴에 직접적인 자산이 됩니다. 관광지의 현장을 아는 개발자가 만드는 서비스이기에 현실적인 문제 해결이 가능합니다.')


add_para('■ KPI 달성 계획', bold=True, size=11)

create_table_light(
    ['KPI', '목표', '달성 방법'],
    [
        ['상품서비스 개시 2건', '웹 서비스 + 모바일 앱', '웹 프로토타입 고도화 + M1~M3 모바일 앱 개발 → M3 정식 출시'],
        ['아이템 검증 20회', '인터뷰+테스트투어+설문', 'M2~M5 제주 현지 테스트투어, 사용자 인터뷰'],
        ['고객확보 5,000명', '누적 가입자', '인플루언서 시딩 10명 + 커뮤니티 바이럴 + 여름 성수기 집중'],
        ['홍보마케팅 25건', '인플루언서+광고+PR', '마케터 채용(M3~), 월 5건 이상 콘텐츠 집행'],
        ['매출 7백만원', '구독(₩3,900)+B2B+광고', 'M3 출시 후 5개월, 저가 구독으로 전환율 8~10% 목표'],
        ['신규고용 4명', '기획PM+개발 2+마케터 1', 'M1부터 순차 채용, 원티드/로켓펀치 활용'],
    ]
)


add_para('■ 대표자 네트워크 및 파트너십', bold=True, size=11)

add_para('• 제주 지역 IT 기업(오션스솔루션) 및 지자체(서귀포시, JDC) 협업 경험 → 관광 현장 네트워크')
add_para('• 크리에이티브힐 Web3 개발팀과의 기술 교류 네트워크 (블록체인 자문)')
add_para('• 법률·세무 전문가 네트워크 확보 완료')


add_para('■ 관광벤처 지원사업의 필요성', bold=True, size=11)

add_para('1. 사업화 자금: 핵심 기능 프로토타입은 구현했으나, 모바일 앱 개발·서비스 고도화·마케팅에 전문 인력이 필요합니다. 지원금으로 기획PM·앱 개발자·블록체인 개발자·마케터를 채용하여 7개월 내 정식 출시와 PMF 검증까지 달성할 수 있습니다.')
add_para('2. 관광 네트워크: 관광벤처 지원사업을 통해 한국관광공사, 지자체 관광공사와의 제휴 기회를 확보할 수 있습니다. 지역 한정 NFT 발행 대행, 관광 데이터 제공 등 B2B 사업의 핵심 파트너를 연결받을 수 있습니다.')
add_para('3. 관광벤처 자격: 관광벤처로 인증받으면 관광기업 육성펀드 투자유치 연계, TIPS 프로그램 연계 등 후속 성장 기회를 확보할 수 있습니다.')
add_para('4. 멘토링: 관광 산업 전문 멘토링을 통해 마케팅 역량 부족(대표 경험 없음)을 보완하고, 관광 시장 진입 전략을 구체화할 수 있습니다.')


doc.add_page_break()


# ===================================================
# 별첨
# ===================================================

add_heading_styled('별첨1: 개인정보 수집·이용 및 제3자 제공 동의서', level=1)
add_para('→ 양식대로 서명. 동의 체크.')

add_heading_styled('별첨2: 창·폐업 및 창업지원사업 수혜이력', level=1)
add_para('• 창·폐업 이력: 없음 (예비창업자)')
add_para('• 창업지원사업 수혜이력: 없음')

add_heading_styled('별첨3: 가점서류', level=1)
add_para("• 청년우대 가점 해당: 1988.11.22생 → 만 37세 → 39세 이하 ('85.02.04 이후 출생)")
add_para('• 주민등록등본 첨부 (공고일 이후 발급, 뒷자리 미포함)')

add_heading_styled('별첨4: 참고자료 (A4 2페이지 이내)', level=1)
add_para('1. 서비스 URL: https://various-belva-untab-1a59bee2.koyeb.app')
add_para('2. MVP 스크린샷 6장:')
add_para('   ① 랜딩 페이지')
add_para('   ② 회원가입/로그인 (카카오/구글)')
add_para('   ③ 여행 그룹 목록')
add_para('   ④ 실시간 채팅')
add_para('   ⑤ NFT 수집 (GPS 인증)')
add_para('   ⑥ AI 매칭 추천')
add_para('3. GitHub: https://github.com/araeLaver/down')


# ── 저장 ──
output_path = os.path.join(os.path.dirname(__file__), '관광벤처_사업계획서_Fryndo_v4.docx')
doc.save(output_path)
print(f'저장 완료: {output_path}')
